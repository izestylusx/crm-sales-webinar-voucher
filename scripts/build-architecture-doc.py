from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "CRM-Architecture-Vision-MVP.docx"

# standard_business_brief preset tokens
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF5EE"
PALE_GOLD = "FFF6D9"
BLACK = "000000"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9DEE7", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def add_numbering(document, fmt, text):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.extend([start, num_fmt, level_text, suffix])

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    p_pr.append(num_pr)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_bullet(doc, text, num_id):
    paragraph = doc.add_paragraph()
    apply_num(paragraph, num_id)
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_numbered(doc, text, num_id):
    return add_bullet(doc, text, num_id)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    set_run_font(paragraph.add_run(f"{label}: "), color=INK, bold=True)
    set_run_font(paragraph.add_run(text), color=INK)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(text), size=9.5, color=INK, bold=True)
    for values in rows:
        row = table.add_row()
        for index, text in enumerate(values):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            set_run_font(paragraph.add_run(str(text)), size=9.5)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header.paragraphs[0]
        header.paragraph_format.space_after = Pt(0)
        set_run_font(header.add_run("Webinar-first MVP | Architecture Vision"), size=9, color=MUTED, bold=True)

        footer = section.footer.paragraphs[0]
        footer.paragraph_format.space_after = Pt(0)
        add_page_number(footer)


def add_cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(42)
    paragraph.paragraph_format.space_after = Pt(4)
    set_run_font(paragraph.add_run("ARCHITECTURE VISION"), size=12, color=BLUE, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    set_run_font(paragraph.add_run("CRM Webinar Management"), size=28, color=INK, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(24)
    set_run_font(paragraph.add_run("Webinar-first MVP berbasis Go untuk salesperson"), size=14, color=MUTED)

    add_table(
        doc,
        ["Dokumen", "Keterangan"],
        [
            ("Audience", "Product, engineering, sales operations, dan leadership"),
            ("Primary user", "Salesperson"),
            ("Active scope", "Booking, reminder, attendance, dan follow-up"),
            ("Backend", "Go modular monolith: crm-api + crm-worker + PostgreSQL"),
            ("Version", "0.3 - 3 September 2026"),
        ],
        [2200, 7160],
    )

    add_callout(
        doc,
        "Keputusan utama",
        "Voucher management, payment integration, subscription activation, school/BOS opportunity, dan commission dipindahkan ke post-MVP.",
        fill=PALE_GREEN,
    )
    add_body(doc, "MVP selesai ketika salesperson dapat mempublikasikan webinar, menerima booking, mengirim reminder, mencatat attendance, dan menyelesaikan follow-up.")
    doc.add_page_break()


def build():
    doc = Document()
    configure_document(doc)
    bullet_id = add_numbering(doc, "bullet", "\u2022")
    number_id = add_numbering(doc, "decimal", "%1.")

    add_cover(doc)

    add_heading(doc, "1. Executive Direction", 1)
    add_callout(doc, "North star", "Satu workspace webinar untuk pekerjaan salesperson, dari publish session sampai follow-up peserta.")
    add_body(doc, "Platform AI pendidikan dan billing tetap berjalan terpisah. Webinar-first MVP tidak membutuhkan payment, provisioning, atau voucher untuk memberikan nilai kepada tim sales.")

    add_heading(doc, "1.1 Outcome MVP", 2)
    for item in [
        "Salesperson membuat dan membagikan webinar dari satu workspace.",
        "Calon client memilih session yang tersedia melalui public booking page.",
        "Capacity, duplicate registration, cancel, dan reschedule ditangani konsisten.",
        "Confirmation, reminder, attendance, dan follow-up memiliki status yang terlihat.",
        "Data attendance siap menjadi future hook tanpa membangun voucher sekarang.",
    ]:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "1.2 Scope boundary", 2)
    add_table(
        doc,
        ["Termasuk MVP", "Dipending ke post-MVP"],
        [
            ("Webinar event dan session", "Voucher lifecycle dan redemption"),
            ("Public booking dan capacity", "Checkout, payment, invoice, refund"),
            ("Confirmation dan reminder", "Subscription activation dan provisioning"),
            ("Attendance manual/CSV", "School opportunity dan dana BOS"),
            ("Participant export dan follow-up", "Commission dan revenue attribution"),
        ],
        [4680, 4680],
    )

    doc.add_page_break()
    add_heading(doc, "2. Users and Webinar Journey", 1)
    add_table(
        doc,
        ["Persona", "MVP responsibility"],
        [
            ("Salesperson", "Create/publish session, invite participant, monitor attendance, follow-up"),
            ("Sales manager", "Team visibility, owner reassignment, attendance/follow-up reporting"),
            ("Marketing/host", "Prepare webinar topic, content, host, and meeting link"),
            ("CRM admin", "Role, template, configuration, and audit access"),
            ("Calon client", "Book, cancel, or reschedule through a secure public flow"),
        ],
        [2400, 6960],
    )

    add_heading(doc, "2.1 Main journey", 2)
    for step in [
        "Salesperson creates a webinar event and publishes one or more sessions.",
        "Calon client opens the public page and chooses an available session.",
        "CRM performs atomic capacity and duplicate checks, then confirms booking.",
        "Worker sends confirmation and scheduled reminders.",
        "Salesperson records attendance manually or through CSV import.",
        "CRM presents attended/no-show participants for follow-up.",
        "Salesperson records follow-up status, note, due date, and outcome.",
    ]:
        add_numbered(doc, step, number_id)

    add_heading(doc, "2.2 State model", 2)
    add_table(
        doc,
        ["Object", "State"],
        [
            ("Session", "draft -> published -> full/completed; published/full -> cancelled"),
            ("Registration", "confirmed -> attended/no_show/cancelled/rescheduled"),
            ("Notification", "pending -> sent/failed/cancelled"),
            ("Follow-up", "not_started -> planned -> contacted -> closed"),
        ],
        [2400, 6960],
    )

    doc.add_page_break()
    add_heading(doc, "3. System Boundaries", 1)
    add_table(
        doc,
        ["Context", "Owner", "MVP responsibility"],
        [
            ("Webinar", "CRM", "Event, session, host, capacity, meeting URL, status"),
            ("Registration", "CRM", "Participant, booking, cancel, reschedule"),
            ("Attendance", "CRM", "Attended/no-show, source, actor, correction audit"),
            ("Follow-up", "CRM", "Owner, task, due date, note, outcome"),
            ("Notification", "CRM", "Confirmation/reminder job and delivery status"),
            ("Identity/learning", "Existing platform", "Outside MVP; user, organization, entitlement, learning access"),
            ("Billing/payment", "Existing billing", "Outside MVP; invoice, payment, settlement, refund"),
        ],
        [1900, 1900, 5560],
    )
    add_callout(doc, "Guardrail", "Booking dan attendance tidak memanggil platform pendidikan atau payment system. Tidak ada shared database.", fill=PALE_GOLD)

    add_heading(doc, "3.1 Logical architecture", 2)
    add_table(
        doc,
        ["Go CRM", "External dependency", "Post-MVP only"],
        [
            ("Sales workspace\nPublic booking\nWebinar/session\nRegistration\nAttendance\nFollow-up", "PostgreSQL\nNotification provider\nExternal meeting link", "Voucher\nPayment/billing integration\nPlatform provisioning\nSchool/BOS pipeline\nCommission"),
        ],
        [3300, 2760, 3300],
    )

    doc.add_page_break()
    add_heading(doc, "4. Domain and Data", 1)
    add_table(
        doc,
        ["Entity", "Purpose", "Key invariant"],
        [
            ("webinar_event", "Reusable webinar topic/template", "Owns one or more sessions"),
            ("webinar_session", "Concrete schedule and capacity", "Only published sessions accept booking"),
            ("participant", "Minimum contact data", "Normalized channel identifier"),
            ("webinar_registration", "Participant booking", "No duplicate active registration per policy"),
            ("attendance_record", "Attended/no-show result", "Correction is audited"),
            ("notification_job", "Confirmation/reminder delivery", "Claimed idempotently by worker"),
            ("follow_up_task", "Sales action after webinar", "Owned and status-tracked"),
            ("audit_log", "Sensitive change trace", "Append-only operational history"),
        ],
        [2550, 3410, 3400],
    )

    add_heading(doc, "4.1 Capacity transaction", 2)
    add_body(doc, "Capacity check, registration insert, confirmation job, and audit record are committed atomically. Concurrent requests cannot exceed available seats. Reschedule reserves the target before releasing the original seat.")

    add_heading(doc, "4.2 Future-compatible data", 2)
    for item in [
        "Use stable registration_id and participant_id.",
        "Store salesperson_id, source, and optional campaign_reference.",
        "Keep attendance source and timestamp auditable.",
        "Do not create voucher, payment, opportunity, or commission tables yet.",
    ]:
        add_bullet(doc, item, bullet_id)

    doc.add_page_break()
    add_heading(doc, "5. API and Background Work", 1)
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ("GET /v1/public/webinars/{token}", "Show published sessions and availability"),
            ("POST /v1/public/webinars/{token}/registrations", "Create booking with capacity/duplicate checks"),
            ("POST /v1/public/registrations/{token}/cancel", "Cancel a registration and pending reminders"),
            ("POST /v1/public/registrations/{token}/reschedule", "Atomically move booking to another session"),
            ("POST /v1/webinar-registrations/{id}/attendance", "Record or correct attendance"),
            ("POST /v1/webinar-sessions/{id}/attendance-imports", "Preview and commit CSV attendance"),
            ("POST /v1/webinar-registrations/{id}/follow-ups", "Create salesperson follow-up"),
            ("GET /v1/webinar-sessions/{id}/export.csv", "Export participant list"),
        ],
        [4900, 4460],
    )

    add_heading(doc, "5.1 Worker jobs", 2)
    for item in [
        "Confirmation and reminder delivery run outside the booking request.",
        "PostgreSQL job rows use locking/lease, bounded concurrency, and retry budget.",
        "Cancellation removes pending reminders idempotently.",
        "Booking remains valid when notification delivery temporarily fails.",
        "Outbound webhook is optional and added only when a real consumer exists.",
    ]:
        add_bullet(doc, item, bullet_id)

    doc.add_page_break()
    add_heading(doc, "6. Go Implementation", 1)
    add_table(
        doc,
        ["Layer", "Go layout", "Responsibility"],
        [
            ("Process", "cmd/crm-api", "Internal API and public booking"),
            ("Process", "cmd/crm-worker", "Confirmation, reminder, retry, cleanup"),
            ("Domain", "internal/webinar", "Event, session, publish, cancel, capacity"),
            ("Domain", "internal/registration", "Booking, duplicate, cancel, reschedule"),
            ("Domain", "internal/attendance", "Manual/CSV attendance and correction"),
            ("Domain", "internal/followup", "Task, owner, status, note, outcome"),
            ("Adapter", "persistence/notification", "PostgreSQL and notification provider"),
        ],
        [1600, 2700, 5060],
    )

    add_heading(doc, "6.1 Runtime rules", 2)
    for item in [
        "Use net/http or a thin router; avoid a heavy framework without team need.",
        "Propagate context.Context to queries and outbound calls.",
        "All goroutines have cancellation, owner, and concurrency bounds.",
        "Never call external providers inside a database transaction.",
        "Support structured logs, metrics, health checks, and graceful shutdown.",
        "Do not add a message broker or microservices for the MVP.",
    ]:
        add_bullet(doc, item, bullet_id)

    doc.add_page_break()
    add_heading(doc, "7. Security, Reliability, and Testing", 1)
    add_table(
        doc,
        ["Area", "MVP control"],
        [
            ("Internal auth", "SSO/OIDC when available, RBAC, team scope"),
            ("Public token", "Opaque random token; no sequential database ID"),
            ("Booking abuse", "Validation, payload limit, rate limit, honeypot/CAPTCHA if needed"),
            ("Sensitive links", "Hash management token; redact meeting URL and PII from logs"),
            ("Capacity", "Transaction + locking/constraint + concurrency tests"),
            ("Reminder", "Durable job, retry budget, failure visibility"),
            ("Attendance", "Preview CSV, row errors, audit correction"),
        ],
        [2500, 6860],
    )

    add_heading(doc, "7.1 Test strategy", 2)
    for item in [
        "Unit: state transition, duplicate policy, reminder schedule, attendance correction.",
        "Integration: capacity locking, reschedule transaction, unique constraint, job claiming.",
        "API: auth/team scope, public token, validation, rate limit, stable error code.",
        "Worker: retry, cancellation, lease expiry, and duplicate execution.",
        "End-to-end: publish -> booking -> reminder -> attendance -> follow-up.",
    ]:
        add_bullet(doc, item, bullet_id)

    doc.add_page_break()
    add_heading(doc, "8. Delivery Roadmap", 1)
    add_table(
        doc,
        ["Phase", "Outcome"],
        [
            ("0. Scope + foundation", "Policy decisions, Go skeleton, migration, CI, OpenAPI"),
            ("1. Webinar/session", "Auth, event/session CRUD, publish/cancel, dashboard"),
            ("2. Public booking", "Capacity, duplicate, confirmation state, cancel/reschedule"),
            ("3. Operate webinar", "Worker, reminders, attendance, follow-up, export"),
            ("4. Hardening", "Concurrency/load tests, runbooks, retention, end-to-end validation"),
            ("Post-MVP", "Provider callback, webhook, voucher, payment, school/BOS, commission"),
        ],
        [2500, 6860],
    )

    add_heading(doc, "8.1 Decisions to lock", 2)
    decisions = [
        "One public page with multiple sessions or one link per session.",
        "Required booking fields and duplicate identifier.",
        "Initial notification channel and reminder schedule.",
        "Manual only or manual plus CSV attendance.",
        "Self-service reschedule deadline.",
        "Team visibility and participant retention policy.",
    ]
    for decision in decisions:
        add_numbered(doc, decision, number_id)

    add_heading(doc, "8.2 Definition of done", 2)
    for item in [
        "No overbooking or duplicate registration under concurrent requests.",
        "Confirmation/reminder status and failure are visible.",
        "Attendance correction and owner reassignment are auditable.",
        "Salesperson can finish follow-up and export participants.",
        "No runtime dependency on voucher, payment, billing, subscription, or provisioning.",
    ]:
        add_bullet(doc, item, bullet_id)

    add_callout(doc, "Recommended next step", "Lock booking fields, duplicate policy, notification channel, attendance method, and team scope before implementation starts.", fill=PALE_GREEN)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

